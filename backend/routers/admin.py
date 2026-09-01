"""
routers/admin.py
Admin-only APIs for users, assignments, and reports.
"""

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from fastapi.responses import Response
import base64
import io
import os
import secrets
import smtplib
import logging
from email.mime.text import MIMEText
from typing import List, Optional
from pydantic import BaseModel, EmailStr

import database as db
from auth_service import require_role, hash_password
from config import settings
from schemas import (
    AdminAssignmentResponse,
    AdminCreateAssignmentRequest,
    AdminReportSummaryResponse,
    AuthUserResponse,
)
from services.ai_service import build_interview_plan_with_usage

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/admin", tags=["admin"])

VIDEO_UPLOAD_DIR = "/home/neosoft/Documents/AiInterview/techscreen/videos"
AUDIO_UPLOAD_DIR = "/home/neosoft/Documents/AiInterview/techscreen/audios"


class AssignmentSummary(BaseModel):
    id: str
    interview_title: str
    target_question_count: Optional[int] = None
    recommended_interview_time_minutes: Optional[int] = None
    status: str
    assigned_at: str


class UserWithAssignments(BaseModel):
    id: str
    name: str
    email: str
    role: str
    assignments: List[AssignmentSummary]


def _plan_fields(item: dict) -> dict:
    """Return interview_title/target_question_count/recommended_interview_time_minutes from assignment plan."""
    plan = item.get("plan") or {}
    return {
        "interview_title": plan.get("interview_title") or "Interview Assignment",
        "target_question_count": plan.get("target_question_count"),
        "recommended_interview_time_minutes": plan.get("recommended_interview_time_minutes"),
    }


def _serialize_user(user: dict) -> AuthUserResponse:
    return AuthUserResponse(
        id=user["id"],
        name=user["name"],
        email=user["email"],
        role=user["role"],
    )


@router.get("/users", response_model=list[AuthUserResponse])
def list_users(_admin=Depends(require_role("admin"))):
    users = db.list_users(role="user")
    return [_serialize_user(user) for user in users]


@router.get("/users-with-assignments", response_model=list[UserWithAssignments])
def list_users_with_assignments(_admin=Depends(require_role("admin"))):
    users = db.list_users(role="user")
    result = []
    for user in users:
        assignments = db.list_assignments_for_user(user["id"])
        user_assignments = [
            AssignmentSummary(
                id=a["id"],
                **_plan_fields(a),
                status=a["status"],
                assigned_at=a["assigned_at"],
            )
            for a in assignments
        ]
        result.append(UserWithAssignments(
            id=user["id"],
            name=user["name"],
            email=user["email"],
            role=user["role"],
            assignments=user_assignments,
        ))
    return result


@router.delete("/assignments/cleanup")
def cleanup_all_assignments(_admin=Depends(require_role("admin"))):
    """Delete all assignments - use for cleanup only."""
    deleted = db.delete_all_assignments()
    return {"deleted": deleted, "message": f"Deleted {deleted} assignment(s)"}


MAX_RESUME_UPLOAD_BYTES = 10 * 1024 * 1024  # 10 MB


def _extract_pdf_text(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    return "\n".join(pages).strip()


def _extract_docx_text(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts).strip()


@router.post("/parse-resume")
async def parse_resume(file: UploadFile = File(...), _admin=Depends(require_role("admin"))):
    """Extract plain text from an uploaded resume file (PDF / DOCX / TXT)."""
    filename = file.filename or "resume"
    ext = os.path.splitext(filename)[1].lower()

    data = await file.read()
    if not data:
        raise HTTPException(400, detail="Uploaded file is empty.")
    if len(data) > MAX_RESUME_UPLOAD_BYTES:
        raise HTTPException(413, detail="File too large (limit 10 MB).")

    try:
        if ext == ".pdf":
            text = _extract_pdf_text(data)
        elif ext in (".docx",):
            text = _extract_docx_text(data)
        elif ext in (".txt",):
            text = data.decode("utf-8", errors="replace").strip()
        elif ext == ".doc":
            raise HTTPException(415, detail="Legacy .doc files are not supported. Please upload .docx, .pdf, or .txt.")
        else:
            raise HTTPException(415, detail=f"Unsupported file type: {ext or 'unknown'}. Use .pdf, .docx, or .txt.")
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(422, detail=f"Could not extract text: {exc}") from exc

    if not text:
        raise HTTPException(422, detail="No text could be extracted (file may be scanned/image-only).")

    return {"filename": filename, "text": text, "char_count": len(text)}


@router.post("/assignments", response_model=AdminAssignmentResponse)
def create_assignment(body: AdminCreateAssignmentRequest, admin=Depends(require_role("admin"))):
    user = db.get_user(body.user_id)
    if not user or user["role"] != "user":
        raise HTTPException(404, detail="User not found")

    if db.has_pending_assignment(body.user_id):
        raise HTTPException(400, detail=f"A pending interview is already assigned to {user['name']}")

    plan, _ = build_interview_plan_with_usage(
        job_description=body.job_description,
        resume_text=body.resume_text,
        years_of_experience=body.years_of_experience,
    )
    assignment = db.create_assignment(
        user_id=body.user_id,
        job_description=body.job_description,
        resume_text=body.resume_text,
        years_of_experience=body.years_of_experience,
        plan=plan,
        assigned_by=admin["id"],
    )
    return AdminAssignmentResponse(
        id=assignment["id"],
        user_id=user["id"],
        user_name=user["name"],
        user_email=user["email"],
        interview_title=plan["interview_title"],
        target_question_count=plan["target_question_count"],
        recommended_interview_time_minutes=plan["recommended_interview_time_minutes"],
        status=assignment["status"],
        assigned_at=assignment["assigned_at"],
        started_at=assignment["started_at"],
        completed_at=assignment["completed_at"],
        session_id=assignment["session_id"],
    )


@router.get("/assignments", response_model=list[AdminAssignmentResponse])
def list_assignments(_admin=Depends(require_role("admin"))):
    assignments = db.list_assignments_with_users()
    return [
        AdminAssignmentResponse(
            id=item["id"],
            user_id=item["user_id"],
            user_name=item["user_name"],
            user_email=item["user_email"],
            **_plan_fields(item),
            status=item["status"],
            assigned_at=item["assigned_at"],
            started_at=item["started_at"],
            completed_at=item["completed_at"],
            session_id=item["session_id"],
        )
        for item in assignments
    ]


@router.get("/reports", response_model=list[AdminReportSummaryResponse])
def list_reports(_admin=Depends(require_role("admin"))):
    reports = db.list_completed_reports()
    return [
        AdminReportSummaryResponse(
            assignment_id=item["assignment_id"],
            session_id=item["session_id"],
            user_id=item["user_id"],
            user_name=item["user_name"],
            user_email=item["user_email"],
            interview_title=item["topic"],
            status=item["status"],
            overall_score=item["overall_score"],
            overall_verdict=item["overall_verdict"],
            questions_answered=item["questions_answered"],
            total_questions=item["total_questions"],
            flag_count=item["flag_count"],
            created_at=item["created_at"],
            completed_at=item["completed_at"],
        )
        for item in reports
    ]


VIDEO_UPLOAD_DIR = "/home/neosoft/Documents/AiInterview/techscreen/videos"


@router.get("/reports/{assignment_id}")
def get_report(assignment_id: str, _admin=Depends(require_role("admin"))):
    assignment = db.get_assignment(assignment_id)
    if not assignment or not assignment.get("session_id"):
        raise HTTPException(404, detail="Report not found")

    session = db.get_session(assignment["session_id"])
    if not session:
        raise HTTPException(404, detail="Session not found")

    user = db.get_user(assignment["user_id"])
    answers = db.get_answers(session["id"])
    flags = db.get_flags(session["id"])

    video_filename = f"{session['id']}.webm"
    video_path = os.path.join(VIDEO_UPLOAD_DIR, video_filename)
    video_available = os.path.exists(video_path)

    scored = [
        answer for answer in answers
        if answer["evaluation"] and isinstance(answer["evaluation"].get("score"), (int, float))
    ]
    avg_score = round((sum(item["evaluation"]["score"] for item in scored) / len(scored)) * 2, 1) if scored else 0.0

    return {
        "assignment": assignment,
        "candidate": {
            "id": user["id"] if user else assignment["user_id"],
            "name": user["name"] if user else "Unknown user",
            "email": user["email"] if user else "",
        },
        "session": session,
        "answers": answers,
        "flags": flags,
        "video_available": video_available,
        "summary": {
            "avg_score": avg_score,
            "questions_answered": len(scored),
            "total_questions": len(session["questions"]),
            "flag_count": len(flags),
        },
    }


@router.delete("/reports/{assignment_id}")
def delete_report(assignment_id: str, _admin=Depends(require_role("admin"))):
    assignment = db.get_assignment(assignment_id)
    if not assignment:
        raise HTTPException(404, detail="Report not found")

    session_id = assignment.get("session_id")
    if session_id and db.get_session(session_id):
        db.delete_session(session_id)

    db.delete_assignment(assignment_id)
    return {"ok": True}


@router.get("/snapshots/{snapshot_id}")
def get_snapshot(snapshot_id: int, _admin=Depends(require_role("admin"))):
    snapshot = db.get_snapshot(snapshot_id)
    if not snapshot:
        raise HTTPException(404, detail="Snapshot not found")
    return Response(content=base64.b64decode(snapshot["image_data"]), media_type="image/jpeg")


# ── Approved candidates (Resume Parser integration) ──────────────────────────

class UpdateEmailRequest(BaseModel):
    email: EmailStr


def _send_credentials_email(to_email: str, name: str, password: str) -> bool:
    """Send login credentials to candidate. Returns True on success."""
    body = (
        f"Hello {name or 'Candidate'},\n\n"
        f"Your interview account has been created.\n\n"
        f"Login URL: {settings.interview_login_url}\n"
        f"Email: {to_email}\n"
        f"Password: {password}\n\n"
        f"Please login and complete your interview.\n\n"
        f"Best regards,\nHR Team\n"
    )
    msg = MIMEText(body)
    msg["Subject"] = "Your Interview Account Credentials"
    msg["From"] = settings.smtp_from_email
    msg["To"] = to_email

    try:
        with smtplib.SMTP(settings.smtp_host, settings.smtp_port, timeout=10) as server:
            if settings.smtp_use_tls:
                server.starttls()
            if settings.smtp_user:
                server.login(settings.smtp_user, settings.smtp_password)
            server.sendmail(settings.smtp_from_email, [to_email], msg.as_string())
        return True
    except Exception as exc:
        logger.warning("Failed to send credentials email to %s: %s", to_email, exc)
        return False


@router.get("/approved-candidates")
def list_approved_candidates(_admin=Depends(require_role("admin"))):
    """Admin sees all candidates approved by HR in Resume Parser."""
    return db.get_approved_candidates()


@router.patch("/approved-candidates/{candidate_id}/email")
def patch_candidate_email(
    candidate_id: str,
    body: UpdateEmailRequest,
    _admin=Depends(require_role("admin")),
):
    """Admin fills in missing email; sets is_ready = TRUE."""
    candidate = db.get_approved_candidate(candidate_id)
    if not candidate:
        raise HTTPException(404, detail="Approved candidate not found")
    db.update_candidate_email(candidate_id, body.email)
    return {"success": True, "message": "Email updated, candidate is now ready"}


@router.post("/approved-candidates/{candidate_id}/assign")
def assign_interview_from_approved(
    candidate_id: str,
    admin=Depends(require_role("admin")),
):
    """Create user account + interview assignment for an approved candidate, then email credentials."""
    candidate = db.get_approved_candidate(candidate_id)
    if not candidate:
        raise HTTPException(404, detail="Approved candidate not found")

    if not candidate["is_ready"] or not candidate["email"]:
        raise HTTPException(400, detail="Candidate email is missing. Please update email first.")

    if db.get_user_by_email(candidate["email"]):
        raise HTTPException(400, detail="User account already exists for this email")

    raw_password = secrets.token_urlsafe(10)

    new_user = db.create_user(
        name=candidate["candidate_name"] or "Candidate",
        email=candidate["email"],
        password_hash=hash_password(raw_password),
        role="user",
    )

    plan, _ = build_interview_plan_with_usage(
        job_description=candidate.get("jd_title") or "Interview",
        resume_text=candidate.get("resume_filename") or "",
        years_of_experience=0,
    )
    assignment = db.create_assignment(
        user_id=new_user["id"],
        job_description=candidate.get("jd_title") or "",
        resume_text="",
        years_of_experience=0,
        plan=plan,
        assigned_by=admin["id"],
    )

    email_sent = _send_credentials_email(
        to_email=candidate["email"],
        name=candidate["candidate_name"] or "Candidate",
        password=raw_password,
    )

    return {
        "success": True,
        "user_id": new_user["id"],
        "assignment_id": assignment["id"],
        "interview_title": plan.get("interview_title"),
        "target_question_count": plan.get("target_question_count"),
        "recommended_interview_time_minutes": plan.get("recommended_interview_time_minutes"),
        "email_sent": email_sent,
        "message": (
            f"Account created and email sent to {candidate['email']}"
            if email_sent
            else f"Account created for {candidate['email']}, but email delivery failed — share credentials manually."
        ),
    }
